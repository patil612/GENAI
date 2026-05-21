from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Literal
from services.ai_service import ai_service
from services.file_parser import extract_text_from_file
import shutil
import os
import uuid
import io
from datetime import datetime

router = APIRouter()

# ── Upload directory ────────────────────────────────────────────────────────────
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ── Supported formats (for validation & UI hints) ───────────────────────────────
SUPPORTED_EXTENSIONS = {
    # Documents
    ".pdf", ".docx", ".doc", ".odt", ".rtf",
    # Plain text / markup
    ".txt", ".md", ".html", ".htm", ".xml", ".json", ".yaml", ".yml",
    ".log", ".ini", ".cfg",
    # Spreadsheets
    ".xlsx", ".xls", ".csv",
    # Presentations
    ".pptx", ".ppt",
    # Images (OCR)
    ".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp", ".gif",
}

MAX_FILE_SIZE_MB = 50  # hard limit per file


# ── Pydantic models ─────────────────────────────────────────────────────────────
class AnalyzeRequest(BaseModel):
    contract_text: str
    policy_text: str


class ClauseRewriteResponse(BaseModel):
    id: int
    original_clause: str
    violated_policy: str
    rewritten_clause: str
    explanation: str
    risk_score: float


class UploadResponse(BaseModel):
    filename: str
    original_name: str
    file_type: str
    size_bytes: int
    extracted_text: str
    char_count: int
    word_count: int
    success: bool
    message: str


# ── Helper ──────────────────────────────────────────────────────────────────────
def _save_upload(file: UploadFile) -> str:
    """Save the uploaded file with a unique name and return the saved path."""
    ext = os.path.splitext(file.filename)[1].lower()
    unique_name = f"{uuid.uuid4().hex}{ext}"
    dest = os.path.join(UPLOAD_DIR, unique_name)
    with open(dest, "wb") as buf:
        shutil.copyfileobj(file.file, buf)
    return dest, unique_name


# ── Routes ──────────────────────────────────────────────────────────────────────

@router.post("/upload", response_model=UploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """
    Accept ANY document or file format and extract its text content.

    Supported formats:
      Documents  : PDF, DOCX, DOC, ODT, RTF
      Plain text : TXT, MD, HTML, XML, JSON, YAML, LOG, INI, CFG
      Spreadsheets: XLSX, XLS, CSV
      Presentations: PPTX, PPT
      Images (OCR): PNG, JPG, JPEG, BMP, TIFF, WEBP, GIF
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided.")

    ext = os.path.splitext(file.filename)[1].lower()

    # ── Extension check ──────────────────────────────────────────────────────
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=(
                f"Unsupported file type '{ext}'. "
                f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            ),
        )

    # ── Save file ────────────────────────────────────────────────────────────
    try:
        saved_path, saved_name = _save_upload(file)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save file: {e}")

    # ── Size check ───────────────────────────────────────────────────────────
    size_bytes = os.path.getsize(saved_path)
    if size_bytes > MAX_FILE_SIZE_MB * 1024 * 1024:
        os.remove(saved_path)
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum allowed size is {MAX_FILE_SIZE_MB} MB.",
        )

    # ── Extract text ─────────────────────────────────────────────────────────
    extracted_text = ""
    success = True
    message = "Text extracted successfully."

    try:
        extracted_text = extract_text_from_file(saved_path, file.filename)
        if not extracted_text or not extracted_text.strip():
            extracted_text = ""
            success = False
            message = "File was processed but no readable text could be extracted."
    except ValueError as ve:
        success = False
        message = str(ve)
        extracted_text = ""
    except Exception as e:
        success = False
        message = f"Unexpected error during text extraction: {e}"
        extracted_text = ""

    word_count = len(extracted_text.split()) if extracted_text else 0

    return UploadResponse(
        filename=saved_name,
        original_name=file.filename,
        file_type=ext.lstrip(".").upper(),
        size_bytes=size_bytes,
        extracted_text=extracted_text,
        char_count=len(extracted_text),
        word_count=word_count,
        success=success,
        message=message,
    )


@router.post("/upload/batch")
async def upload_files_batch(files: List[UploadFile] = File(...)):
    """
    Upload multiple files at once. Returns an array of UploadResponse objects.
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files provided.")
    if len(files) > 10:
        raise HTTPException(status_code=400, detail="Maximum 10 files per batch upload.")

    results = []
    for file in files:
        try:
            # Reuse single-upload logic via internal call
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in SUPPORTED_EXTENSIONS:
                results.append({
                    "original_name": file.filename,
                    "success": False,
                    "message": f"Unsupported file type '{ext}'.",
                    "extracted_text": "",
                })
                continue

            saved_path, saved_name = _save_upload(file)
            size_bytes = os.path.getsize(saved_path)

            extracted_text = ""
            success = True
            message = "Text extracted successfully."
            try:
                extracted_text = extract_text_from_file(saved_path, file.filename)
                if not extracted_text.strip():
                    success = False
                    message = "No readable text could be extracted."
                    extracted_text = ""
            except Exception as e:
                success = False
                message = str(e)
                extracted_text = ""

            results.append({
                "filename": saved_name,
                "original_name": file.filename,
                "file_type": ext.lstrip(".").upper(),
                "size_bytes": size_bytes,
                "extracted_text": extracted_text,
                "char_count": len(extracted_text),
                "word_count": len(extracted_text.split()) if extracted_text else 0,
                "success": success,
                "message": message,
            })
        except Exception as e:
            results.append({
                "original_name": file.filename,
                "success": False,
                "message": f"Failed: {e}",
                "extracted_text": "",
            })

    return results


@router.get("/upload/supported-formats")
def get_supported_formats():
    """Returns the list of all supported file formats grouped by category."""
    return {
        "documents": [".pdf", ".docx", ".doc", ".odt", ".rtf"],
        "plain_text": [".txt", ".md", ".html", ".htm", ".xml", ".json", ".yaml", ".yml", ".log", ".ini", ".cfg"],
        "spreadsheets": [".xlsx", ".xls", ".csv"],
        "presentations": [".pptx", ".ppt"],
        "images_ocr": [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp", ".gif"],
        "max_file_size_mb": MAX_FILE_SIZE_MB,
    }


@router.post("/analyze", response_model=List[ClauseRewriteResponse])
async def analyze_contract(request: AnalyzeRequest):
    if not request.contract_text or not request.policy_text:
        raise HTTPException(status_code=400, detail="Contract text and policy text are required.")

    results = ai_service.analyze_contract(request.contract_text, request.policy_text)
    return results


# ── Download Models ──────────────────────────────────────────────────────────────
class DownloadClause(BaseModel):
    id: int
    original_clause: str
    violated_policy: str
    rewritten_clause: str
    explanation: str
    risk_score: float

class DownloadRequest(BaseModel):
    results: List[DownloadClause]
    contract_name: Optional[str] = "Contract"
    format: Literal["docx", "txt"] = "docx"


# ── Download Routes ──────────────────────────────────────────────────────────────

@router.post("/download/docx")
async def download_docx(request: DownloadRequest):
    """
    Generate a formatted DOCX redline report from analysis results and stream it for download.
    """
    try:
        import docx
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx is required. Run: pip install python-docx")

    doc = docx.Document()

    # ── Page margins ────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Title ────────────────────────────────────────────────────────────────────
    title = doc.add_heading("Auto-Redline Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)  # slate-900

    # ── Metadata block ───────────────────────────────────────────────────────────
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Document: {request.contract_name}   |   "
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}   |   "
        f"Violations Found: {len(request.results)}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)  # slate-500
    doc.add_paragraph()

    # ── Summary table ────────────────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["#", "Risk Score", "Violated Policy", "Status"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True

    for item in request.results:
        row = table.add_row().cells
        row[0].text = str(item.id + 1)
        row[1].text = f"{int(item.risk_score * 100)}/100"
        row[2].text = item.violated_policy[:80] + ("…" if len(item.violated_policy) > 80 else "")
        row[3].text = "⚠ Non-Compliant"

    doc.add_paragraph()

    # ── Detailed violations ───────────────────────────────────────────────────────
    doc.add_heading("Detailed Violations & Redlines", level=1)

    for i, item in enumerate(request.results, 1):
        # Violation heading
        h = doc.add_heading(f"Violation {i}  —  Risk Score: {int(item.risk_score * 100)}/100", level=2)
        h.runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)  # red-600

        # Violated policy
        doc.add_heading("Violated Policy", level=3)
        p = doc.add_paragraph(item.violated_policy)
        p.paragraph_format.left_indent = Inches(0.3)

        # Original clause (struck-through in red)
        doc.add_heading("Original Clause (Non-Compliant)", level=3)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(item.original_clause)
        run.font.strike = True
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)  # red

        # Rewritten clause (bold green)
        doc.add_heading("AI Rewritten Clause (Compliant)", level=3)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(item.rewritten_clause)
        run.bold = True
        run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)  # emerald-600

        # Explanation
        doc.add_heading("GenAI Explanation", level=3)
        p = doc.add_paragraph(item.explanation)
        p.paragraph_format.left_indent = Inches(0.3)

        # Separator
        doc.add_paragraph("─" * 60)

    # ── Footer ────────────────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Generated by Auto-Redline · Powered by Llama 3 GenAI")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # ── Stream ────────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = "".join(c for c in request.contract_name if c.isalnum() or c in " _-").strip()
    filename = f"Auto-Redline_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/download/txt")
async def download_txt(request: DownloadRequest):
    """
    Generate a plain-text redline report from analysis results and stream it for download.
    """
    lines = [
        "=" * 70,
        "  AUTO-REDLINE ANALYSIS REPORT",
        f"  Document : {request.contract_name}",
        f"  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"  Violations Found: {len(request.results)}",
        "=" * 70,
        "",
    ]

    for i, item in enumerate(request.results, 1):
        lines += [
            f"VIOLATION {i}  |  Risk Score: {int(item.risk_score * 100)}/100",
            "-" * 70,
            f"Violated Policy:",
            f"  {item.violated_policy}",
            "",
            f"Original Clause (Non-Compliant):",
            f"  [STRUCK] {item.original_clause}",
            "",
            f"AI Rewritten Clause (Compliant):",
            f"  {item.rewritten_clause}",
            "",
            f"GenAI Explanation:",
            f"  {item.explanation}",
            "",
            "=" * 70,
            "",
        ]

    lines.append("Generated by Auto-Redline · Powered by Llama 3 GenAI")
    content = "\n".join(lines)

    buf = io.BytesIO(content.encode("utf-8"))
    safe_name = "".join(c for c in request.contract_name if c.isalnum() or c in " _-").strip()
    filename = f"Auto-Redline_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

    return StreamingResponse(
        buf,
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/health")
def health_check():
    return {"status": "healthy", "supported_formats": len(SUPPORTED_EXTENSIONS)}
